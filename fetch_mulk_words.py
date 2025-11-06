# requirements:
# requests
# pandas
# openpyxl
# python-docx
# reportlab
# arabic-reshaper
# python-bidi

import argparse
import os
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse, parse_qs

import pandas as pd
import requests
from docx import Document
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

try:
    import arabic_reshaper  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    arabic_reshaper = None

try:
    from bidi.algorithm import get_display  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    get_display = None


CONTENT_API_BASE_URL = "https://apis.quran.foundation"
API_URL = f"{CONTENT_API_BASE_URL}/content/api/v4/verses/by_chapter/67"
OAUTH2_BASE_URL = "https://oauth2.quran.foundation"
OAUTH_TOKEN_URL = f"{OAUTH2_BASE_URL}/oauth2/token"
DOCS_URL = "https://api-docs.quran.foundation/docs/content_apis_versioned/verses-by-chapter-number"
TRANSLATION_ID = 85

AyahData = Tuple[int, List[Dict[str, str]], str, str]


FONT_CANDIDATE_NAMES = [
    "arialuni.ttf",
    "Arial Unicode.ttf",
    "Arial Unicode MS.ttf",
    "arial.ttf",
    "Tahoma.ttf",
    "trado.ttf",
    "Times New Roman.ttf",
    "Amiri-Regular.ttf",
    "Amiri.ttf",
    "NotoNaskhArabic-Regular.ttf",
    "NotoSansArabic-Regular.ttf",
    "Scheherazade-Regular.ttf",
    "DejaVuSans.ttf",
]

FONT_SEARCH_DIRS = [
    Path("C:/Windows/Fonts"),
    Path("/System/Library/Fonts"),
    Path("/Library/Fonts"),
    Path("/usr/share/fonts"),
    Path("/usr/local/share/fonts"),
    Path.home() / ".fonts",
    Path.home() / ".local" / "share" / "fonts",
]


def _locate_arabic_font() -> Optional[Path]:
    """Find a TrueType font that can render Arabic glyphs."""
    for directory in FONT_SEARCH_DIRS:
        if not directory.exists():
            continue
        for candidate in FONT_CANDIDATE_NAMES:
            candidate_path = directory / candidate
            if candidate_path.exists():
                return candidate_path
    return None


def register_arabic_font() -> str:
    """Register and return a font name for Arabic rendering in PDFs."""
    font_path = _locate_arabic_font()
    if font_path:
        font_name = f"ArabicFont-{font_path.stem}"
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            return font_name
        except Exception:  # pragma: no cover - fallback in case registration fails
            pass
    return "Helvetica"


ARABIC_FONT_NAME = register_arabic_font()


def shape_arabic_text(text: str) -> str:
    """Perform basic reshaping/bidi handling for Arabic text if libraries are available."""
    if not text:
        return text
    reshaped = text
    if arabic_reshaper is not None:
        try:
            reshaped = arabic_reshaper.reshape(text)
        except Exception:
            reshaped = text
    if get_display is not None:
        try:
            reshaped = get_display(reshaped)
        except Exception:
            pass
    return reshaped


class MissingEnvironment(Exception):
    """Raised when required environment variables are missing."""


def require_env_var(name: str) -> str:
    """Fetch an environment variable or raise an informative error."""
    value = os.getenv(name)
    if not value:
        raise MissingEnvironment(
            f"Environment variable {name} is required. Kindly export it before running this script."
        )
    return value


def fetch_page(
    session: requests.Session,
    url: str,
    headers: Dict[str, str],
    params: Dict[str, str],
    retries: int = 3,
    backoff_seconds: float = 1.5,
) -> Dict:
    """Fetch a single API page with simple retry/backoff for transient failures."""
    attempt = 0
    while True:
        try:
            response = session.get(
                url,
                headers=headers,
                params=params,
                timeout=15,
            )
            if 500 <= response.status_code < 600:
                raise requests.HTTPError(
                    f"Server error {response.status_code} on {response.url}"
                )
            response.raise_for_status()
            return response.json()
        except (requests.RequestException, ValueError) as exc:
            attempt += 1
            if attempt > retries:
                raise
            sleep_for = backoff_seconds * attempt
            time.sleep(sleep_for)


def fetch_access_token(
    session: requests.Session,
    client_id: str,
    client_secret: str,
    retries: int = 3,
    backoff_seconds: float = 1.5,
) -> str:
    """Obtain an OAuth2 client-credentials access token."""
    payload = "grant_type=client_credentials&scope=content"
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    attempt = 0
    while True:
        try:
            response = session.post(
                OAUTH_TOKEN_URL,
                auth=(client_id, client_secret),
                headers=headers,
                data=payload,
                timeout=15,
            )
            if 500 <= response.status_code < 600:
                raise requests.HTTPError(
                    f"OAuth server error {response.status_code} on {response.url}"
                )
            response.raise_for_status()
            data = response.json()
        except (requests.RequestException, ValueError) as exc:
            attempt += 1
            if attempt > retries:
                raise
            sleep_for = backoff_seconds * attempt
            time.sleep(sleep_for)
            continue

        token = data.get("access_token")
        if token:
            return token
        error_description = data.get("error_description") or data.get("error")
        raise SystemExit(
            f"Failed to obtain access token. Details: {error_description or data}"
        )


def extract_page_number(link: str) -> Optional[int]:
    """Extract an integer page number from a pagination link."""
    if not link:
        return None
    try:
        query = urlparse(link).query
        params = parse_qs(query)
        if "page" in params and params["page"]:
            return int(params["page"][0])
    except (ValueError, TypeError):
        return None
    return None


def paginate_all(
    session: requests.Session,
    headers: Dict[str, str],
    per_page: int,
    sleep_between_pages: float,
    word_translation_language: str,
) -> List[Dict]:
    """Retrieve all verses by following pagination hints until exhausted."""
    verses: List[Dict] = []
    page = 1
    while True:
        params = {
            "words": "true",
            "word_fields": "text_uthmani",
            "per_page": str(per_page),
            "page": str(page),
            "word_translation_language": word_translation_language,
            "translations": str(TRANSLATION_ID),
            "translation_fields": "text",
        }
        data = fetch_page(session, API_URL, headers, params)

        if not isinstance(data, dict):
            print("Unexpected response type; expected JSON object.", file=sys.stderr)
            raise SystemExit(f"See API docs: {DOCS_URL}")

        verses_batch = data.get("verses")
        if verses_batch is None:
            print(
                "Response missing 'verses' key. Available keys:",
                list(data.keys()),
                file=sys.stderr,
            )
            raise SystemExit(f"See API docs: {DOCS_URL}")

        if not isinstance(verses_batch, list):
            print("'verses' key is not a list.", file=sys.stderr)
            raise SystemExit(f"See API docs: {DOCS_URL}")

        if not verses_batch:
            break

        verses.extend(verses_batch)

        links = data.get("links") or {}
        pagination_meta = data.get("meta", {}).get("pagination", {})

        next_page = extract_page_number(links.get("next"))
        if next_page:
            page = next_page
        else:
            current_page = pagination_meta.get("current_page")
            total_pages = pagination_meta.get("total_pages")
            meta_next_page = pagination_meta.get("next_page")
            if meta_next_page:
                page = meta_next_page
            elif current_page and total_pages and current_page < total_pages:
                page = current_page + 1
            else:
                break

        if sleep_between_pages > 0:
            time.sleep(sleep_between_pages)

    return verses


def parse_ayah_number(verse: Dict) -> Optional[int]:
    """Derive an ayah number from supported keys."""
    number = verse.get("verse_number")
    if isinstance(number, int):
        return number

    verse_key = verse.get("verse_key")
    if isinstance(verse_key, str) and ":" in verse_key:
        try:
            return int(verse_key.split(":", maxsplit=1)[1])
        except ValueError:
            return None
    return None


def normalize_verse(verse: Dict) -> AyahData:
    """Convert raw verse JSON into structured records."""
    ayah_number = parse_ayah_number(verse)
    if ayah_number is None:
        raise ValueError("Unable to determine ayah number for verse.")

    words = verse.get("words") or []
    word_records: List[Dict[str, str]] = []
    tokens: List[str] = []
    verse_translation_text = ""

    translations = verse.get("translations") or []
    if isinstance(translations, list):
        for translation in translations:
            if not isinstance(translation, dict):
                continue
            translation_id = translation.get("resource_id") or translation.get("id")
            if str(translation_id) == str(TRANSLATION_ID):
                verse_translation_text = translation.get("text") or ""
                break
        if not verse_translation_text and translations:
            fallback = translations[0]
            if isinstance(fallback, dict):
                verse_translation_text = fallback.get("text") or ""

    for idx, word in enumerate(words, start=1):
        text_uthmani = ""
        translation_text = ""
        if isinstance(word, dict):
            text_uthmani = word.get("text_uthmani") or ""
            word_translation = word.get("word_translation") or word.get("translation") or {}
            if isinstance(word_translation, dict):
                translation_text = word_translation.get("text") or ""
            elif isinstance(word.get("translations"), list):
                translations_list = word.get("translations")
                if translations_list:
                    first_translation = translations_list[0]
                    if isinstance(first_translation, dict):
                        translation_text = first_translation.get("text") or ""
        word_records.append(
            {
                "Word #": idx,
                "Uthmani": text_uthmani,
                "Translation": translation_text,
            }
        )
        if text_uthmani:
            tokens.append(text_uthmani.strip())

    ayah_text = " ".join(tokens)
    return ayah_number, word_records, ayah_text, verse_translation_text


def export_excel(ayah_data: List[AyahData], path: str) -> None:
    """Export ayah-wise word data into an Excel workbook."""
    ayah_data_sorted = sorted(ayah_data, key=lambda item: item[0])
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for ayah_number, word_records, _, _ in ayah_data_sorted:
            sheet_name = f"Ayah {ayah_number}"
            if word_records:
                df = pd.DataFrame(
                    word_records, columns=["Word #", "Uthmani", "Translation"]
                )
            else:
                df = pd.DataFrame(columns=["Word #", "Uthmani", "Translation"])
            df.to_excel(writer, sheet_name=sheet_name, index=False)

            worksheet = writer.sheets[sheet_name]
            for idx, column in enumerate(df.columns, start=1):
                series = df[column].astype(str)
                max_length = max([len(column)] + [len(value) for value in series])
                adjusted = min(max(12, max_length + 2), 60)
                worksheet.column_dimensions[get_column_letter(idx)].width = adjusted


def export_docx(ayah_data: List[AyahData], path: str) -> None:
    """Export ayah-wise data into a Word document with headings and tables."""
    document = Document()
    document.add_heading("Surah Al-Mulk (67) — Word by Word", level=0)

    for ayah_number, word_records, ayah_text, verse_translation in sorted(
        ayah_data, key=lambda item: item[0]
    ):
        document.add_heading(f"Ayah {ayah_number}", level=1)
        document.add_paragraph(ayah_text if ayah_text else "(No text available)")
        document.add_paragraph(
            verse_translation if verse_translation else "(Translation unavailable)"
        )

        table = document.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Word #"
        header_cells[1].text = "Uthmani"
        header_cells[2].text = "Translation"

        for record in word_records:
            row = table.add_row().cells
            row[0].text = str(record["Word #"])
            row[1].text = record["Uthmani"]
            row[2].text = record.get("Translation", "")

        # Add a blank paragraph to separate tables visually.
        document.add_paragraph()

    document.save(path)


def _build_pdf_flowables(
    ayah_number: int,
    ayah_text: str,
    verse_translation: str,
    word_records: List[Dict[str, str]],
) -> List[Any]:
    """Construct flowables representing a single ayah page."""
    styles = getSampleStyleSheet()
    heading_style = styles["Heading1"]
    body_style = styles["BodyText"]
    translation_style = ParagraphStyle(
        "Translation",
        parent=body_style,
        leading=16,
        spaceAfter=12,
    )
    arabic_style = ParagraphStyle(
        "Arabic",
        parent=body_style,
        fontName=ARABIC_FONT_NAME,
        leading=18,
        spaceAfter=10,
        alignment=TA_RIGHT,
    )

    ayah_text_display = (
        shape_arabic_text(ayah_text) if ayah_text else "(No text available)"
    )
    flowables: List[Any] = [
        Paragraph(f"Ayah {ayah_number}", heading_style),
        Spacer(1, 0.2 * inch),
        Paragraph(ayah_text_display, arabic_style),
        Spacer(1, 0.1 * inch),
        Paragraph(
            verse_translation if verse_translation else "(Translation unavailable)",
            translation_style,
        ),
        Spacer(1, 0.2 * inch),
    ]

    table_data = [["Word #", "Arabic", "English"]]
    for record in word_records:
        table_data.append(
            [
                str(record["Word #"]),
                shape_arabic_text(record["Uthmani"]),
                record.get("Translation", ""),
            ]
        )

    table = Table(table_data, colWidths=[1 * inch, 2.5 * inch, 3 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 11),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (1, 1), (1, -1), ARABIC_FONT_NAME),
                ("ALIGN", (1, 1), (1, -1), "RIGHT"),
                ("FONTNAME", (0, 1), (0, -1), "Helvetica"),
            ]
        )
    )
    flowables.append(table)

    return flowables


def export_pdf_combined(ayah_data: List[AyahData], path: str) -> None:
    """Create a single PDF with one page per ayah."""
    doc = SimpleDocTemplate(path, pagesize=A4)
    story: List[Any] = []
    for ayah_number, word_records, ayah_text, verse_translation in sorted(
        ayah_data, key=lambda item: item[0]
    ):
        story.extend(
            _build_pdf_flowables(
                ayah_number, ayah_text, verse_translation, word_records
            )
        )
        story.append(PageBreak())

    if story and isinstance(story[-1], PageBreak):
        story.pop()

    doc.build(story)


def export_pdf_single_pages(ayah_data: List[AyahData], directory: Path) -> None:
    """Create individual single-page PDFs for each ayah."""
    directory.mkdir(parents=True, exist_ok=True)
    for ayah_number, word_records, ayah_text, verse_translation in sorted(
        ayah_data, key=lambda item: item[0]
    ):
        output_path = directory / f"ayah_{ayah_number:02d}.pdf"
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        flowables = _build_pdf_flowables(
            ayah_number, ayah_text, verse_translation, word_records
        )
        doc.build(flowables)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Download Surah Al-Mulk word-by-word data and export to Excel and Word."
    )
    parser.add_argument("--per-page", type=int, default=50, help="Items per request page.")
    parser.add_argument(
        "--sleep",
        type=float,
        default=0.0,
        help="Seconds to sleep between page requests.",
    )
    parser.add_argument("--excel", default="surah_mulk_words.xlsx", help="Excel output path.")
    parser.add_argument("--docx", default="surah_mulk_words.docx", help="Word document output path.")
    parser.add_argument(
        "--translation-language",
        default="en",
        help="ISO code or language id for word-by-word translations (default: en).",
    )
    parser.add_argument("--pdf", default="surah_mulk_words.pdf", help="Combined PDF output path.")
    parser.add_argument(
        "--pdf-dir",
        default="surah_mulk_single_pages",
        help="Directory to store single-page PDFs (one per ayah).",
    )
    args = parser.parse_args()

    with requests.Session() as session:
        try:
            client_id = require_env_var("QF_CLIENT_ID")
            client_secret = require_env_var("QF_CLIENT_SECRET")
        except MissingEnvironment as exc:
            print(exc, file=sys.stderr)
            raise SystemExit(1)

        try:
            access_token = fetch_access_token(
                session=session,
                client_id=client_id,
                client_secret=client_secret,
            )
        except requests.RequestException as exc:
            print(f"Token request failed: {exc}", file=sys.stderr)
            raise SystemExit("Unable to fetch OAuth token.")

        headers = {
            "Accept": "application/json",
            "x-auth-token": access_token,
            "x-client-id": client_id,
        }

        sessions_verses = paginate_all(
            session=session,
            headers=headers,
            per_page=args.per_page,
            sleep_between_pages=args.sleep,
            word_translation_language=args.translation_language,
        )

    ayah_data: List[AyahData] = []
    for verse in sessions_verses:
        try:
            ayah_number, word_records, ayah_text, verse_translation = normalize_verse(
                verse
            )
        except ValueError as exc:
            print(
                f"Skipping verse due to parsing issue: {exc}. Raw keys: {list(verse.keys())}",
                file=sys.stderr,
            )
            continue
        ayah_data.append((ayah_number, word_records, ayah_text, verse_translation))

    ayah_data.sort(key=lambda item: item[0])

    if len(ayah_data) != 30:
        print(
            f"Warning: expected 30 ayahs but collected {len(ayah_data)}.",
            file=sys.stderr,
        )

    export_excel(ayah_data, args.excel)
    export_docx(ayah_data, args.docx)
    export_pdf_combined(ayah_data, args.pdf)
    export_pdf_single_pages(ayah_data, Path(args.pdf_dir))

    total_words = sum(len(records) for _, records, _, _ in ayah_data)
    excel_path = os.path.abspath(args.excel)
    docx_path = os.path.abspath(args.docx)
    pdf_path = os.path.abspath(args.pdf)
    pdf_dir = os.path.abspath(args.pdf_dir)

    print(f"Verses fetched: {len(ayah_data)}")
    print(f"Total words: {total_words}")
    print(f"Excel output: {excel_path}")
    print(f"Word output: {docx_path}")
    print(f"Combined PDF output: {pdf_path}")
    print(f"Single-page PDFs directory: {pdf_dir}")


if __name__ == "__main__":
    try:
        main()
    except requests.RequestException as exc:
        print(f"Request failed: {exc}", file=sys.stderr)
        raise SystemExit(f"See API docs: {DOCS_URL}")
